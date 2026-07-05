# -*- coding: utf-8 -*-
"""
Genera los documentos Word del soporte documental a partir de los .md de la
carpeta 'soporte documental/'. Cada Markdown se convierte a un .docx con
portada institucional (logo CORPOURABA), estilo verde, tablas y pie de página.

Uso:  python etl/generar_soporte_docx.py
"""
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

ROOT = Path(__file__).resolve().parent.parent
CARPETA = ROOT / "soporte documental"
LOGO = ROOT / "frontend" / "public" / "logo-corpouraba.png"

VERDE_OSC = RGBColor(0x0B, 0x3D, 0x25)
VERDE = RGBColor(0x1F, 0x73, 0x47)
GRIS = RGBColor(0x6B, 0x72, 0x80)
NEGRO = RGBColor(0x22, 0x2A, 0x25)
FUENTE = "Calibri"


# ── formato en línea (**negrita**, *cursiva*, `código`, [texto](url)) ─────────
def _pre_links(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\((https?://[^)]+)\)", r"\1 (\2)", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", text)
    return text


def add_inline(par, text: str):
    text = _pre_links(text)
    for tk in re.split(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)", text):
        if not tk:
            continue
        if tk.startswith("**") and tk.endswith("**"):
            r = par.add_run(tk[2:-2]); r.bold = True
        elif tk.startswith("`") and tk.endswith("`"):
            r = par.add_run(tk[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x9F, 0x12, 0x39)
        elif tk.startswith("*") and tk.endswith("*") and len(tk) > 2:
            r = par.add_run(tk[1:-1]); r.italic = True
        else:
            par.add_run(tk)


def heading(doc, text, level):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    sizes = {1: 17, 2: 13.5, 3: 11.5, 4: 10.5}
    r = p.add_run(text.strip())
    r.bold = True
    r.font.name = FUENTE
    r.font.size = Pt(sizes.get(level, 11))
    r.font.color.rgb = VERDE_OSC if level <= 2 else VERDE
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(4)
    if level == 1:
        _borde_inferior(p)


def _borde_inferior(p):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "1F7347")
    pbdr.append(bottom)
    pPr.append(pbdr)


def tabla(doc, filas):
    if not filas:
        return
    ncol = max(len(f) for f in filas)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, fila in enumerate(filas):
        celdas = t.add_row().cells
        for j in range(ncol):
            txt = fila[j] if j < len(fila) else ""
            celda = celdas[j]
            celda.paragraphs[0].text = ""
            add_inline(celda.paragraphs[0], txt)
            for run in celda.paragraphs[0].runs:
                run.font.size = Pt(9.5)
                if i == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            celda.paragraphs[0].paragraph_format.space_after = Pt(2)
            if i == 0:
                _sombrear(celda, "1F7347")


def _sombrear(celda, hexcolor):
    tcPr = celda._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def bullet(doc, text, nivel=0, numerada=False):
    estilo = "List Number" if numerada else "List Bullet"
    try:
        p = doc.add_paragraph(style=estilo)
    except KeyError:
        p = doc.add_paragraph()
        p.add_run("• ")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * nivel)
    p.paragraph_format.space_after = Pt(2)
    add_inline(p, text)
    for r in p.runs:
        if r.font.size is None:
            r.font.size = Pt(10.5)
    return p


def parrafo(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_inline(p, text)
    for r in p.runs:
        if r.font.size is None:
            r.font.size = Pt(10.5)
        r.font.color.rgb = NEGRO if r.font.color.rgb is None else r.font.color.rgb


# ── portada y pie ────────────────────────────────────────────────────────────
def portada(doc, titulo):
    if LOGO.exists():
        pl = doc.add_paragraph()
        pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
        try:
            pl.add_run().add_picture(str(LOGO), width=Inches(1.9))
        except Exception:
            pass
    sup = doc.add_paragraph()
    r = sup.add_run("OBSERVATORIO DE DEFORESTACIÓN CORPOURABA")
    r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = GRIS
    sup.paragraph_format.space_before = Pt(10)
    sup.paragraph_format.space_after = Pt(2)

    pt = doc.add_paragraph()
    rt = pt.add_run(titulo)
    rt.bold = True; rt.font.size = Pt(24); rt.font.color.rgb = VERDE_OSC
    pt.paragraph_format.space_after = Pt(4)
    _borde_inferior(pt)

    sub = doc.add_paragraph()
    rs = sub.add_run(
        "CORPOURABA — Corporación para el Desarrollo Sostenible del Urabá\n"
        "Soporte documental de la plataforma web · Deforestación 2000–2024\n"
        "Plataforma creada por Alberto Vivas y Carlos Zuluaga"
    )
    rs.font.size = Pt(10); rs.font.color.rgb = GRIS
    sub.paragraph_format.space_after = Pt(10)


def pie(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.text = ""
    r = p.add_run("Soporte documental · Observatorio de Deforestación CORPOURABA")
    r.font.size = Pt(8); r.font.color.rgb = GRIS
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ── conversión Markdown → docx ───────────────────────────────────────────────
def md_a_docx(md: str, salida: Path):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FUENTE
    normal.font.size = Pt(10.5)
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.9)
        s.top_margin = s.bottom_margin = Inches(0.8)

    lineas = md.replace("\r\n", "\n").split("\n")
    # título = primer H1
    titulo = "Documento"
    idx0 = 0
    for i, ln in enumerate(lineas):
        if ln.startswith("# "):
            titulo = ln[2:].strip()
            idx0 = i + 1
            break
    portada(doc, titulo)
    pie(doc)

    i = idx0
    n = len(lineas)
    while i < n:
        ln = lineas[i]
        s = ln.strip()
        # tabla
        if s.startswith("|") and i + 1 < n and re.match(r"^\|?[\s:|-]+\|?$", lineas[i + 1].strip()) and "-" in lineas[i + 1]:
            filas = []
            while i < n and lineas[i].strip().startswith("|"):
                fila_txt = lineas[i].strip()
                if re.match(r"^\|?[\s:|-]+\|?$", fila_txt) and "-" in fila_txt:
                    i += 1
                    continue
                celdas = [c.strip() for c in fila_txt.strip("|").split("|")]
                filas.append(celdas)
                i += 1
            tabla(doc, filas)
            continue
        if not s:
            i += 1
            continue
        m = re.match(r"^(#{1,4})\s+(.*)$", s)
        if m:
            heading(doc, m.group(2), len(m.group(1)))
            i += 1
            continue
        if re.match(r"^-{3,}$", s) or re.match(r"^\*{3,}$", s):
            i += 1
            continue
        mnum = re.match(r"^(\s*)\d+\.\s+(.*)$", ln)
        if mnum:
            nivel = len(mnum.group(1)) // 2
            bullet(doc, mnum.group(2), nivel, numerada=True)
            i += 1
            continue
        mbul = re.match(r"^(\s*)[-*+]\s+(.*)$", ln)
        if mbul:
            nivel = len(mbul.group(1)) // 2
            bullet(doc, mbul.group(2), nivel, numerada=False)
            i += 1
            continue
        parrafo(doc, s)
        i += 1

    doc.save(str(salida))
    return salida


def main():
    if not CARPETA.exists():
        print(f"No existe la carpeta {CARPETA}")
        return 1
    mds = sorted(CARPETA.glob("*.md"))
    if not mds:
        print("No hay .md que convertir.")
        return 1
    for md_path in mds:
        md = md_path.read_text(encoding="utf-8")
        salida = md_path.with_suffix(".docx")
        md_a_docx(md, salida)
        print(f"OK  {salida.name}  ({salida.stat().st_size // 1024} KB)")
    print(f"\n{len(mds)} documentos generados en '{CARPETA.name}'.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
